#python module
import os
from datetime import date, timedelta
import json
from io import StringIO, BytesIO
import logging
#django Module
from django.shortcuts import render
from django.views.generic import TemplateView, CreateView, ListView, DetailView, UpdateView
from django.contrib.auth.mixins import LoginRequiredMixin, UserPassesTestMixin
from django.contrib.auth.decorators import login_required
from django.core.exceptions import PermissionDenied
from django.core import serializers
from django.views.decorators.csrf import csrf_exempt
from django.http import JsonResponse
from django.http import HttpResponse, HttpResponseRedirect
from django.urls import reverse
from django.contrib import messages
from django.conf import settings

from django.db.models import Q, F
from django.db.models import Count, Sum, Max

#3rd party module
# from rest_framework import viewsets
# from rest_framework import permissions
from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework import status

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from openpyxl import load_workbook

#My module
from .models import HomeRequest, CoResident
from .forms import HomeRequestForm, CoResidentFormSet
from apps.Utility.Constants import (YEARROUND_PROCESSSTEP, HomeRequestProcessStep,PERSON_STATUS, 
                                    HomeRentPermission, FINANCE_CODE)
from apps.UserData.models import Unit
from apps.Configurations.models import YearRound
from apps.UserData.forms import UserCurrentDataForm
from apps.UserData.models import User
from apps.Payment.models import FinanceData
from .serializers import HomeRequestSerializer
logger = logging.getLogger('MainLog')

def get_current_year():
    CurrentYearRound = YearRound.objects.filter(CurrentStep__in = ['RS','UP','PP'])
    CurrentYear = CurrentYearRound[0].Year
    # print('CurrentYear = ',CurrentYear)
    return CurrentYear


class AuthenUserTestMixin(LoginRequiredMixin, UserPassesTestMixin):
    login_url = '/login' 
    allow_groups = []

    def test_func(self):
        for ag in self.allow_groups:
            if self.request.user.groups.filter(name=ag).exists():
                return True
        return False

    def has_home_request(self):
        queryset = HomeRequest.objects.filter(Requester = self.request.user)
        queryset = queryset.filter(year_round__Year = get_current_year())
        return queryset.exists()


class ProcessFlow(TemplateView):
    template_name = "HomeRequest/process_flow.html"

class CreateHomeRequestView(AuthenUserTestMixin, CreateView):
    allow_groups = ['RTAF_NO_HOME_USER']
    model = HomeRequest
    form_class = HomeRequestForm
    template_name = "HomeRequest/CreateHomeRequest.html"

    # ทดสอบเพิ่มเติมว่าถ้าปีนี้มีการส่งคำขอแล้ว ก็ส่งอีกไม่ได้
    def test_func(self):
        if super().test_func() == False:
            return False
        else:
            return not super().has_home_request()

    def get(self, request, *args, **kwargs):
        self.object = None
        # form_class = self.get_form_class()
        # form = self.get_form(form_class)
        initial_value = {
                            'Rank': request.user.Rank,
                            'FullName': request.user.FullName,
                            'Position': request.user.Position,
                            'Unit': request.user.CurrentUnit,
                            'MobilePhone' : request.user.MobilePhone,
                            'OfficePhone' : request.user.OfficePhone,
                            'Unit' : request.user.CurrentUnit,
                            'Salary' : request.user.current_salary,
                            'Status' : request.user.current_status,
                            'SpouseName' : request.user.current_spouse_name,
                            'SpousePID' : request.user.current_spouse_pid,
                            'IsHRISReport' : request.user.current_spouse_pid,
                            'Address' : request.user.Address
                        }       

        # เช็คข้อมูลการเบิก คชบ.ของตนเอง ในช่วง 6 เดือนล่าสุด
        homerent_data = FinanceData.objects.filter(PersonID =  request.user.PersonID
                                          ).filter(date__gte = date.today() - timedelta(days = 185)
                                          ).filter(code = FINANCE_CODE.HOMERENT
                                          ).filter(money__gt = 0
                                          ).order_by("money")
            
        print('homerent_data',homerent_data)
        if(homerent_data.exists()):
            initial_value['RentPermission'] = HomeRentPermission.used
            initial_value['have_rent'] = True
            initial_value['RentalCost'] = homerent_data[0].money
        else:
            initial_value['RentPermission'] = HomeRentPermission.no_permission

        # ถ้ามีสถานภาพสมรส
        if(request.user.current_status in [PERSON_STATUS.MARRIES_TOGETHER, PERSON_STATUS.MARRIES_SEPARATE]):
            # เช็คข้อมูลการเบิก คชบ.ของคู่สมรส ในช่วง 6 เดือนล่าสุด
            homerent_data = FinanceData.objects.filter(PersonID =  request.user.current_spouse_pid
                                          ).filter(date__gte = date.today() - timedelta(days = 185)
                                          ).filter(code = FINANCE_CODE.HOMERENT
                                          ).filter(money__gt = 0
                                          ).order_by("money")
            if(homerent_data.exists()):
                initial_value['have_rent_spouse'] = True
                initial_value['RentalCostSpouse'] = homerent_data[0].money


        form = self.form_class(initial = initial_value, prefix='hr')

        co_resident_formset = CoResidentFormSet()

        initial_value = {
                            'OfficePhone' : request.user.OfficePhone,
                            'MobilePhone' : request.user.MobilePhone,
                            'RTAFEMail': request.user.username + '@rtaf.mi.th',
                        }        
        user_current_data_form =  UserCurrentDataForm(initial = initial_value, prefix='userdata')
        return self.render_to_response(
                                self.get_context_data(form=form,
                                                      co_resident_formset=co_resident_formset,
                                                      user_current_data_form = user_current_data_form
                                                        ))    

    # def get(self, request, *args, **kwargs):

    #     form = self.form_class(initial = initial_value)
    #     return render(request, self.template_name, {'form': form})

    def post(self, request, *args, **kwargs):
        self.object = None
        form = self.form_class(request.POST, request.FILES, prefix='hr')

        user_current_data_form = UserCurrentDataForm(self.request.POST,instance = request.user, prefix='userdata')
        co_resident_formset = CoResidentFormSet(self.request.POST)
        if form.is_valid() and co_resident_formset.is_valid():
            return self.form_valid(form, user_current_data_form, co_resident_formset)
        else:
            return self.form_invalid(form, user_current_data_form, co_resident_formset)

    def form_valid(self, form, user_current_data_form, co_resident_formset):
        self.object = form.save(commit=False)
        # กำหนดค่าเริ่มต้นให้ form    
        self.object.Requester = self.request.user
        CurrentYearRound = YearRound.objects.filter(CurrentStep__in = ['RS','UP','PP'])

        self.object.year_round = CurrentYearRound[0]
        self.object.Unit = self.request.user.CurrentUnit
        self.object.save()
        
        user_current_data_form.save()

        co_resident = co_resident_formset.save(commit=False)
        for cr in co_resident:
            cr.home_request = self.object
            cr.save()

        if 'save' in self.request.POST:            
            self.object.ProcessStep = HomeRequestProcessStep.REQUESTER_PROCESS
            self.object.save()
            messages.success(self.request, f'เพิ่มข้อมูลบ้านพักของ {self.object.FullName} เรียบร้อย')
        elif 'send' in self.request.POST:
            self.object.update_process_step(
                                    HomeRequestProcessStep.REQUESTER_SENDED, 
                                    self.request.user)
            self.object.save()
            messages.success(self.request, f'บันทึกและส่งข้อมูลบ้านพักของ {self.object.FullName} เรียบร้อย')


        return HttpResponseRedirect(self.get_success_url())

    def form_invalid(self, form, user_current_data_form, co_resident_formset):
        # print(form.errors)
        return self.render_to_response(
                 self.get_context_data(form=form,
                                       co_resident_formset=co_resident_formset,
                                       user_current_data_form = user_current_data_form
                                       )
        )


class UpdateHomeRequestView(AuthenUserTestMixin, UpdateView):
    allow_groups = ['RTAF_NO_HOME_USER']
    model = HomeRequest
    form_class = HomeRequestForm
    template_name = "HomeRequest/CreateHomeRequest.html"

    #ในกรณีที่ส่งรายงานแล้ว จะไม่สามารถแก้ไขข้อมูลได้
    def test_func(self):
        if super().test_func() == False:
            return False
        else:
            self.object = self.get_object()
            return not self.object.RequesterSended

    def get(self, request, *args, **kwargs):
        self.object = self.get_object()
        home_request = self.object
        
        co_resident_formset = CoResidentFormSet(instance = home_request, 
                                                queryset = home_request.CoResident.order_by("Relation"))

        # print("UpdateHomeRequestView:get")
        

        return self.render_to_response(
                  self.get_context_data(form = HomeRequestForm(instance=self.object, prefix='hr'),
                                        user_current_data_form =  UserCurrentDataForm(instance = request.user, prefix='userdata'),
                                        co_resident_formset = co_resident_formset
                                        )
                                       )

    def get_context_data(self, **kwargs):
        data = super().get_context_data(**kwargs)
        # print("UpdateHomeRequestView:get_context_data")

        if self.request.GET:
            # print('get_context_data:self.request.GET')
            data['form'] = HomeRequestForm(instance=self.object, prefix='hr')
            data["user_current_data_form"] = UserCurrentDataForm(instance = self.request.user, prefix='userdata')
            data["co_resident"] = CoResidentFormSet(instance=self.object)
        elif self.request.POST:
            # print('get_context_data:self.request.POST')
            data['form'] = HomeRequestForm(self.request.POST, prefix='hr')
            data["user_current_data_form"] = UserCurrentDataForm(self.request.POST, prefix='userdata')
            data["co_resident_formset"] = CoResidentFormSet(self.request.POST)
        return data

    def post(self, request, *args, **kwargs):
        self.object = self.get_object()
        form = self.form_class(request.POST, request.FILES, instance = self.object, prefix='hr')

        user_current_data_form = UserCurrentDataForm(self.request.POST, instance = request.user, prefix='userdata')
        co_resident_formset = CoResidentFormSet(self.request.POST, instance = self.object)

        # print('UpdateHomeRequestView:post:form.is_valid() ',form.is_valid())
        
        # print('UpdateHomeRequestView:post:user_current_data_form.is_valid() ',user_current_data_form.is_valid())
        # print('UpdateHomeRequestView:post:co_resident_formset.is_valid() ',co_resident_formset.is_valid())

        if form.is_valid() and user_current_data_form.is_valid() and co_resident_formset.is_valid():
            # print('co_resident_formset ', co_resident_formset)
            return self.form_valid(form, user_current_data_form, co_resident_formset)
        else:
            return self.form_invalid(form, user_current_data_form, co_resident_formset)

    def form_valid(self, form, user_current_data_form, co_resident_formset):

        # print('UpdateHomeRequestView:form_valid co_resident_formset = ',co_resident_formset)
        user_current_data_form.save()
        co_resident = co_resident_formset.save(commit=False)

        for cr in co_resident_formset.deleted_objects:
            cr.delete()

        for cr in co_resident:
            # print('coresident = ',cr)
            cr.home_request = self.object
            cr.save()

        if 'save' in self.request.POST:
            messages.warning(self.request, f'บันทึกการแก้ไขข้อมูลบ้านพักของ {self.object.FullName} เรียบร้อย *** ขั้นตอนการส่งยังไม่เรียบร้อย ***')
        elif 'send' in self.request.POST:
            self.object.update_process_step(
                                    HomeRequestProcessStep.REQUESTER_SENDED, 
                                    self.request.user)
            self.object.save()
            messages.success(self.request, f'บันทึกการแก้ไขและส่งข้อมูลบ้านพักของ {self.object.FullName} เรียบร้อย')

        return super().form_valid(form)

    def form_invalid(self, form, user_current_data_form, co_resident_formset):
        # print('UpdateHomeRequestView:form_invalid => user_current_data_form  ') #, user_current_data_form)
        # print('UpdateHomeRequestView:form_invalid => co_resident_formset => ') #, co_resident_formset)
        return super().form_invalid(form)

    def get_success_url(self):        
        return reverse('HomeRequest:af_person')


class AFPersonListView(AuthenUserTestMixin,ListView):
    template_name = "HomeRequest/af_person.html"
    allow_groups = ['RTAF_NO_HOME_USER']

    def get_queryset(self):
        queryset = HomeRequest.objects.filter(Requester = self.request.user)
        queryset = queryset.order_by("-year_round__Year")
        return queryset

    def get_context_data(self, **kwargs):        
        context = super().get_context_data(**kwargs)
        context['has_request'] = super().has_home_request(**kwargs)
        return context        


class HomeRequestUnitListView(AuthenUserTestMixin, ListView):
    model = HomeRequest    
    template_name = "HomeRequest/list.html"
    allow_groups = ['PERSON_ADMIN', 'PERSON_UNIT_ADMIN']

    def get_queryset(self, *args, **kwargs):

        if self.request.user.groups.filter(name='PERSON_UNIT_ADMIN').exists():
                queryset = HomeRequest.objects.filter(Unit = self.request.user.CurrentUnit)                
        
        if 'unit_id' in self.kwargs:
            unit_id = self.kwargs['unit_id']
            if self.request.user.groups.filter(name='PERSON_ADMIN').exists():
                    queryset = HomeRequest.objects.filter(Unit_id = unit_id)
        
        queryset = queryset.filter(year_round__Year = get_current_year())
        queryset = queryset.order_by("-year_round__Year")
        return queryset 
    

class HomeRequestAdminListView(HomeRequestUnitListView):
    model = HomeRequest    
    template_name = "Person/modal_list.html"
    allow_groups = ['PERSON_ADMIN']

    def get_queryset(self, *args, **kwargs):
        
        if 'unit_id' in self.kwargs:
            unit_id =  self.kwargs['unit_id']
            if self.request.user.groups.filter(name='PERSON_ADMIN').exists():
                    queryset = HomeRequest.objects.filter(Unit_id = unit_id)
        
        queryset = queryset.filter(year_round__Year = get_current_year())
        queryset = queryset.order_by("-year_round__Year")
        return queryset

class UnitList4PersonAdmin(AuthenUserTestMixin, APIView):
    allow_groups = ['PERSON_ADMIN']

    def get(self, request, *args, **kwargs):        
        unit_id = kwargs["unit_id"] if kwargs.get("unit_id", None) is not None else 41

        queryset = HomeRequest.objects.filter(Unit_id = unit_id).order_by("Rank")
        queryset = queryset.filter(year_round__Year = get_current_year())
        serializer = HomeRequestSerializer(queryset, many=True)
        return Response(serializer.data)


class HomeRequestUnitSummaryListView(AuthenUserTestMixin,ListView):
    # template_name = "Person/unit_summary.html"
    template_name = "Person/unit_summary_vue.html"
    allow_groups = ['PERSON_ADMIN']

    def get_queryset(self):

        Num_RP = Count('ProcessStep', filter = Q(ProcessStep = HomeRequestProcessStep.REQUESTER_PROCESS))
        Num_RS = Count('ProcessStep', filter = Q(ProcessStep = HomeRequestProcessStep.REQUESTER_SENDED))
        Num_UP = Count('ProcessStep', filter = Q(ProcessStep = HomeRequestProcessStep.UNIT_PROCESS))
        Num_US = Count('ProcessStep', filter = Q(ProcessStep = HomeRequestProcessStep.UNIT_SENDED))
        Num_PP = Count('ProcessStep', filter = Q(ProcessStep = HomeRequestProcessStep.PERSON_PROCESS))
        Num_PA = Count('ProcessStep', filter = Q(ProcessStep = HomeRequestProcessStep.PERSON_ACCEPTED))
        Num_GH = Count('ProcessStep', filter = Q(ProcessStep = HomeRequestProcessStep.GET_HOUSE))
        DateSended = Max('UnitDateApproved')
        UnitApprover = Max('UnitApprover__first_name')
        
        queryset = HomeRequest.objects.filter(year_round__Year = get_current_year())

        queryset = queryset.exclude(
                                    Q(ProcessStep = HomeRequestProcessStep.REQUESTER_CANCEL) 
                                ).values('Unit'
                                ).annotate(
                                    DateSended = DateSended,
                                    UnitApprover = UnitApprover,
                                    Num_RP = Num_RP,
                                    Num_RS = Num_RS,
                                    Num_UP = Num_UP,
                                    Num_US = Num_US,
                                    Num_PP = Num_PP,
                                    Num_PA = Num_PA,
                                    Num_GH = Num_GH
                                ).values(
                                    'Unit',
                                    'Unit__ShortName',
                                    'DateSended',
                                    'UnitApprover',
                                    'Num_RP',
                                    'Num_RS',
                                    'Num_UP',
                                    'Num_US',
                                    'Num_PP',
                                    'Num_PA',
                                    'Num_GH'
                                )
        
        return queryset

class HomeRequestDetail(AuthenUserTestMixin, DetailView):
    allow_groups = ['RTAF_NO_HOME_USER','PERSON_UNIT_ADMIN']
    model = HomeRequest
    template_name = "HomeRequest/Detail.html"

    def get_context_data(self, *args, **kwargs):
        context = super(HomeRequestDetail, self).get_context_data(*args, **kwargs)            
        co_residence = CoResident.objects.filter(home_request=self.object).order_by("Relation")
        context["co_residence"] = co_residence
        return context


@login_required
def cancel_request(request, home_request_id):
    home_request = HomeRequest.objects.get(id = home_request_id)
    if request.user != home_request.Requester:
        raise PermissionDenied()
    else:
        if request.method == "POST":
            data = json.loads(request.body.decode("utf-8"))
            Comment = data["comment"]
            # print("data =  " ,Comment)
        home_request.cancel_request = True
        home_request.Comment = home_request.Comment + 'แจ้งยกเลิกคำขอเมื่อ ' + str(date.today()) + Comment
        home_request.save()
        return JsonResponse({"ok": True}, safe=False)
    
               
@login_required
def update_process_step(request, home_request_id, process_step):
    allow_groups = ['PERSON_ADMIN', 'PERSON_UNIT_ADMIN']
    allow_access = False
    for ag in allow_groups:
        if request.user.groups.filter(name=ag).exists():
            allow_access = True
            break
    if not allow_access:
        raise PermissionDenied()
               
    home_request = HomeRequest.objects.get(id = home_request_id)
    home_request.ProcessStep = process_step
    if process_step == 'RP':
        home_request.RequesterDateSend = None
    home_request.save()
    home_request.update_process_step(process_step, request.user)

    if process_step in [ HomeRequestProcessStep.REQUESTER_PROCESS,
                         HomeRequestProcessStep.UNIT_PROCESS,
                         HomeRequestProcessStep.UNIT_SENDED]:
        messages.info(request,f'บันทึกขั้นตอนคำขอบ้าน {home_request.Requester.FullName} เรียบร้อย')
        return HttpResponseRedirect("/hr/list")
    elif process_step in [ HomeRequestProcessStep.PERSON_PROCESS, 
                           HomeRequestProcessStep.PERSON_ACCEPTED]:
        return JsonResponse({"success": True})


@csrf_exempt
def homerequest_detail(request, username):
    try:
        user = User.objects.get(username = username)     
        CurrentYearRound = YearRound.objects.filter(CurrentStep__in = ['RS','UP','PP'])
        year_round = CurrentYearRound[0] 
        homerequest = HomeRequest.objects.filter(year_round = year_round).filter(Requester = user)
    except User.DoesNotExist:
        dump = json.dumps({'status': 'username not found'})            
        return HttpResponse(dump, content_type='application/json')

    print('homerequest',homerequest)
    if not homerequest.exists():
        dump = json.dumps({'status': 'hr not found'})            
        return HttpResponse(dump, content_type='application/json')

    if request.method == 'GET':
        serializer = HomeRequestSerializer(homerequest[0])
        return JsonResponse(serializer.data)


def ArabicToThai(number_string): 
    dic = { 
        '0':'๐', 
        '1':'๑', 
        '2':'๒', 
        '3':'๓', 
        '4':'๔', 
        '5':'๕', 
        '6':'๖', 
        '7':'๗', 
        '8':'๘', 
        '9':'๙', 
    }
    result = ""
    for ch in number_string:
        if ch in "0123456789":
            result += dic[ch]
        else:
            result += ch
    
    return result

month_text = ["","ม.ค.","ก.พ.","มี.ค.","เม.ย.","พ.ค.","มิ.ย.","ก.ค.","ส.ค.","ก.ย.","ต.ค.","พ.ย.","ธ.ค."]
full_month_text = ["","มกราคม","กุมภาพันธ์","มีนาคม","เมษายน","พฤษภาคม","มิถุนายน","กรกฎาคม","สิงหาคม","กันยายน","ตุลาคม","พฤศจิกายน","ธันวาคม"]

def TestDocument(request, home_request_id):
    # testdoc = static('documents/test_doc.docx')
    home_request = HomeRequest.objects.get(id = home_request_id)

    if home_request.ProcessStep == 'RP':
        testdoc =  os.path.join(settings.TEMPLATES[0]['DIRS'][0],'documents/house_request_data_draft.docx')
        docx_title= f"Draft-{home_request.Requester.AFID}.docx"
    else:
        testdoc =  os.path.join(settings.TEMPLATES[0]['DIRS'][0],'documents/house_request_data.docx')
        docx_title= f"House-{home_request.Requester.AFID}.docx"

    document = Document(testdoc)

    self_call = "กระผม" if home_request.Requester.Sex == "ชาย" else "ดิฉัน"
    SpouseName = home_request.SpouseName if home_request.SpouseName else ""

    
    Salary = int(home_request.Salary) if home_request.Salary is not None else "0"
    add_salary = int(home_request.AddSalary) if home_request.AddSalary is not None else "0"

    Income = int(Salary) + int(add_salary)
    
    RentalCost = ""
    if home_request.RentPermission == 1:
        Z1, Z2, Z3 = "X", "  ", "  "
    elif home_request.RentPermission == 2:
        Z1, Z2, Z3 = "  ", "X", "  "
    elif home_request.RentPermission == 3:
        Z1, Z2, Z3 = "  ", "X", "X"
        RentalCost = home_request.RentalCost  if home_request.RentalCost is not None else 0
        RentalCost = "{:,} บาท".format(RentalCost)
    
    if home_request.have_rent_spouse:
        ZK = "X"
        WifeRentalCost = home_request.RentalCostSpouse if home_request.RentalCostSpouse  is not None  else 0
        WifeRentalCost = "{:,} บาท".format(WifeRentalCost)        
    else:
        ZK = "  "
        WifeRentalCost = " -"  

    if home_request.Status in [2, 7]: # ถ้าสถานะเป็นสมรส
        IsNotRTAFHome = "X" if home_request.IsNotRTAFHome else "  "
        SpouseName = "{}".format(SpouseName)
    else:
        IsNotRTAFHome = "-"
        ZK = "-"
        SpouseName = " - "

    IsNotBuyHome = "X" if home_request.IsNotBuyHome else "  "
    IsNotOwnHome = "X" if home_request.IsNotOwnHome else "  "
    IsNeverRTAFHome = "X" if home_request.IsNeverRTAFHome else "  "
    ImportanceDuty = "X" if home_request.ImportanceDuty else "  "
    IsMoveFromOtherUnit = "X" if home_request.IsMoveFromOtherUnit else "  "
    IsHomelessEvict = "X" if home_request.IsHomelessEvict else "  "
    ContinueHouse = "X" if home_request.ContinueHouse else "  "
    IsHomelessDisaster = "X" if home_request.IsHomelessDisaster else "  "
    OtherTrouble = "X" if home_request.OtherTrouble else "  "
    IsHomeNeed = "X" if home_request.IsHomeNeed else "  "
    IsFlatNeed = "X" if home_request.IsFlatNeed else "  "
    IsShopHouseNeed = "X" if home_request.IsShopHouseNeed else "  "
    NumResidence = home_request.CoResident.all().count()
    NumResidence = "-" if NumResidence == 0 else NumResidence

    HouseRegistration = "X" if home_request.HouseRegistration else "  "
    MarriageRegistration = "X" if home_request.MarriageRegistration else "  "
    SpouseApproved = "X" if home_request.SpouseApproved else "  "
    DivorceRegistration = "X" if home_request.DivorceRegistration else "  "
    SpouseDeathRegistration = "X" if home_request.SpouseDeathRegistration else "  "
    Month = month_text[date.today().month]
    Year =  str((date.today().year + 543) % 100)

    PR1 = home_request.get_ZoneRequestPriority1_display() if home_request.ZoneRequestPriority1 else " - "
    PR2 = home_request.get_ZoneRequestPriority2_display() if home_request.ZoneRequestPriority2 else " - "
    PR3 = home_request.get_ZoneRequestPriority3_display() if home_request.ZoneRequestPriority3 else " - "
    PR4 = home_request.get_ZoneRequestPriority4_display() if home_request.ZoneRequestPriority4 else " - "
    PR5 = home_request.get_ZoneRequestPriority5_display() if home_request.ZoneRequestPriority5 else " - "
    PR6 = home_request.get_ZoneRequestPriority6_display() if home_request.ZoneRequestPriority6 else " - "
    Rank = "{}".format(home_request.Requester.get_Rank_display())
    if "ว่าที่" in Rank:
        Rank = Rank[7:]
    FullName = home_request.FullName
    if "ว่าที่" in FullName:
        FullName = FullName[7:]
    dic = {
            'Sex': self_call,
            'Rank': Rank,
            'FullName':FullName,
            'PersonID':home_request.Requester.PersonID,
            'Position':home_request.Position,
            'ADDR':home_request.Address,
            'GPC':home_request.GooglePlusCodes1,
            'Distance': str(home_request.distance),
            'UnitFN':home_request.Unit.FullName,
            'UnitSN':home_request.Unit.ShortName,
            'OfficePhone':home_request.Requester.OfficePhone,
            'MobilePhone':home_request.Requester.MobilePhone,
            'Income': "{:,}".format(Income),
            'Status':"{}".format(home_request.get_Status_display()),
            'SpouseName': SpouseName,
            'Z1':Z1,
            'Z2':Z2,
            'Z3':Z3,
            'WifeRentalCost':WifeRentalCost,
            'RentalCost' : RentalCost,
            'Z4':IsNotBuyHome,
            'Z5':IsNotOwnHome,
            'Z6':IsNeverRTAFHome,
            'Z7':IsNotRTAFHome,
            'Z8':ImportanceDuty,
            'Z9':IsMoveFromOtherUnit,
            'ZA':IsHomelessEvict,
            'ZB':IsHomelessDisaster,
            'ZC':IsHomeNeed,
            'ZD':IsFlatNeed,
            'ZE':IsShopHouseNeed,
            'ZF':HouseRegistration,
            'ZG':MarriageRegistration,
            'ZH':SpouseApproved,
            'ZI':DivorceRegistration,
            'ZJ':SpouseDeathRegistration,
            'ZK':ZK,
            'ZX':ContinueHouse,
            'ZY':OtherTrouble,
            'PR1': PR1,
            'PR2': PR2,
            'PR3': PR3,
            'PR4': PR4,
            'PR5': PR5,
            'PR6': PR6,
            'Residence' : str(NumResidence),
            'Month': Month,
            'Year':Year
            }
    # print(dic)
    
    for para in document.paragraphs:
        print('para = ',para)
        for key, value in dic.items():
            if key in para.text:
                inline = para.runs
                # print('inline = ',inline)
                # Loop added to work with runs (strings with same style)
                for i in range(len(inline)):
                    if key in inline[i].text:
                        if value: # ถ้ามีการกรอกข้อมูล
                            Value = ArabicToThai(value) if key != 'GPC' else value
                        else: # ถ้าไม่มีการกรอกข้อมูล
                            Value = ""
                        text = inline[i].text.replace(key, Value)
                        inline[i].text = text

    # Prepare document for download        
    # -----------------------------
    f = BytesIO()
    document.save(f)
    length = f.tell()
    f.seek(0)
    response = HttpResponse(
        f.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    # print('docx_title = ',docx_title)
    response['Content-Disposition'] = 'attachment; filename="{}"'.format(docx_title)
    response['Content-Length'] = length
    return response

def ConsentForm(request):
    testdoc =  os.path.join(settings.TEMPLATES[0]['DIRS'][0],'documents/letter_of_consent.docx')
    docx_title= f"consentform.docx"
    

    document = Document(testdoc)

    spname = request.POST.get("if_spouse_name","")
    address = request.POST.get("if_address","")

    spouse = "ภรรยา" if request.user.Sex == "ชาย" else "สามี"
    selfcall = "สามี" if request.user.Sex == "ชาย" else "ภรรยา"


    day = str(date.today().day)
    month = full_month_text[date.today().month]
    year =  str((date.today().year + 543))

    dic = {
            'spouse': spouse,            
            'selfcall': selfcall,            
            'spname': spname,            
            'address': address,            
            'username':request.user.FullName,
            'day' : day,
            'month' : month,
            'year' : year,
            }

    for para in document.paragraphs:
        print('para = ',para)
        for key, value in dic.items():
            if key in para.text:
                inline = para.runs
                # print('inline = ',inline)
                # Loop added to work with runs (strings with same style)
                for i in range(len(inline)):
                    if key in inline[i].text:
                        if value: # ถ้ามีการกรอกข้อมูล
                            Value = ArabicToThai(value) if key != 'GPC' else value
                        else: # ถ้าไม่มีการกรอกข้อมูล
                            Value = ""
                        text = inline[i].text.replace(key, Value)
                        inline[i].text = text


    f = BytesIO()
    document.save(f)
    length = f.tell()
    f.seek(0)
    response = HttpResponse(
        f.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    # print('docx_title = ',docx_title)
    response['Content-Disposition'] = 'attachment; filename="{}"'.format(docx_title)
    response['Content-Length'] = length
    return response

 
def TestExcel(request,unit_id):
    # testdoc = static('documents/test_doc.docx')
    testxls =  os.path.join(settings.TEMPLATES[0]['DIRS'][0],'documents/unit_report.xlsx')
    # Start by opening the spreadsheet and selecting the main sheet
    workbook = load_workbook(filename=testxls)
    xls_title= f"unit_report.xlsx"

    sheet = workbook.active

    # Write what you want into a specific cell
    sheet["A1"] = f'หน่วย {request.user.CurrentUnit.ShortName}'

    
    xls_unit = Unit.objects.get(id = unit_id)

    queryset = HomeRequest.objects.filter(Unit = xls_unit)
    
    queryset = queryset.filter(year_round__Year = get_current_year())
    queryset = queryset.order_by("-year_round__Year")
    first_row = 6
    for i, data in enumerate(queryset):
        sheet[f"A{first_row+i}"] = i+1
        sheet[f"B{first_row+i}"] = data.FullName
        sheet[f"C{first_row+i}"] = data.get_Status_display()
        sheet[f"D{first_row+i}"] = data.Salary + data.AddSalary
        sheet[f"R{first_row+i}"] = data.Requester.MobilePhone

    f = BytesIO()
    
    workbook.save(f)
    length = f.tell()
    f.seek(0)
    response = HttpResponse(
        f.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    )
    # print('xls = ',xls_title)
    response['Content-Disposition'] = 'attachment; filename="{}"'.format(xls_title)
    response['Content-Length'] = length
    return response