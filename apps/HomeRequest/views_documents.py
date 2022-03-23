#python module
import os
from datetime import date, timedelta
import json
from io import StringIO, BytesIO
import logging
#django Module
from django.shortcuts import render

from django.contrib.auth.decorators import login_required
from django.views.decorators.csrf import csrf_exempt
from django.http import HttpResponse,  HttpResponseBadRequest, HttpResponseForbidden
from django.http import Http404
from django.core.exceptions import PermissionDenied
from django.urls import reverse
from django.conf import settings
from django.db.models import Count, Q, F
from django.http.response import JsonResponse

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from openpyxl import load_workbook
from copy import copy

from linebot import LineBotApi, WebhookParser
from linebot.exceptions import InvalidSignatureError, LineBotApiError
from linebot.models import MessageEvent, TextMessage, TextSendMessage

line_bot_api = LineBotApi('odFxpwpkdg')
parser = WebhookParser('ab315a0889e13')

#My module
from .models import CoResident, HomeRequest
from .views import get_current_year
from apps.Configurations.models import YearRound
from apps.UserData.models import Unit
from apps.Utility.utils import decryp_file
from apps.Utility.Constants import PERSON_STATUS, officer_rank, non_officer_rank, status_single, status_family
logger = logging.getLogger('MainLog')
evidence_logger = logging.getLogger('EvidenceAccessLog')


def ArabicToThai(number_string): 
    if not number_string:
        return number_string
    if number_string == "0":
        return "-"
    dic = { 
        '0':'๐', 
        '1':'๑', 
        '2':'๒', 
        '3':'๓', 
        '4':'๔', 
        '5':'๕', 
        '6':'๖', 
        '7':'๗', 
        '8':'๘', 
        '9':'๙', 
    }
    result = ""
    for ch in number_string:
        if ch in "0123456789":
            result += dic[ch]
        else:
            result += ch
    
    return result

month_text = ["","ม.ค.","ก.พ.","มี.ค.","เม.ย.","พ.ค.","มิ.ย.","ก.ค.","ส.ค.","ก.ย.","ต.ค.","พ.ย.","ธ.ค."]
full_month_text = ["","มกราคม","กุมภาพันธ์","มีนาคม","เมษายน","พฤษภาคม","มิถุนายน","กรกฎาคม","สิงหาคม","กันยายน","ตุลาคม","พฤศจิกายน","ธันวาคม"]


def TestDocument(request, home_request_id,detail_doc = 0):
    # testdoc = static('documents/test_doc.docx')
    home_request = HomeRequest.objects.get(id = home_request_id)

    # if home_request.ProcessStep == 'RP':
    #     testdoc =  os.path.join(settings.TEMPLATES[0]['DIRS'][0],'documents/house_request_data_1_page_draft.docx')
    #     docx_title= f"Draft-{home_request.Requester.AFID}.docx"
    # else:
    if detail_doc:
        testdoc =  os.path.join(settings.TEMPLATES[0]['DIRS'][0],'documents/house_request_detail.docx')
    else:
        testdoc =  os.path.join(settings.TEMPLATES[0]['DIRS'][0],'documents/house_request_data_1_page.docx')
    docx_title= f"House-{home_request.Requester.AFID}.docx"

    document = Document(testdoc)

    self_call = "กระผม" if home_request.Requester.Sex == "ชาย" else "ดิฉัน"
    SpouseName = home_request.SpouseName if home_request.SpouseName else ""

    
    Salary = int(home_request.Salary) if home_request.Salary is not None else "0"
    add_salary = int(home_request.AddSalary) if home_request.AddSalary is not None else "0"

    Income = int(Salary) + int(add_salary)
    
    RentalCost = ""
    if home_request.RentPermission == 1:
        Z1, Z2, Z3 = "X", "  ", "  "
    elif home_request.RentPermission == 2:
        Z1, Z2, Z3 = "  ", "X", "  "
    elif home_request.RentPermission == 3:
        Z1, Z2, Z3 = "  ", "X", "X"
        RentalCost = home_request.RentalCost  if home_request.RentalCost is not None else 0
        RentalCost = "{:,} บาท".format(RentalCost)
    
    if home_request.have_rent_spouse:
        ZK = "X"
        WifeRentalCost = home_request.RentalCostSpouse if home_request.RentalCostSpouse  is not None  else 0
        WifeRentalCost = "{:,} บาท".format(WifeRentalCost)        
    else:
        ZK = "  "
        WifeRentalCost = " -"  

    IsNotRTAFHome = "--"
    status_description = ""
    if home_request.Status in [2, 7]: # ถ้าสถานะเป็นสมรส
        IsNotRTAFHome = "X" if home_request.IsNotRTAFHome else "  "
        SpouseName = "{}".format(SpouseName)
        status_description = "มีสถานภาพสมรส"
        family_single = "ครอบครัว"
    elif  home_request.Status == 1:
        status_description = "มีสถานภาพโสด"
        family_single = "โสด"
    else:
        ZK = "--"
        SpouseName = " - "
        family_single = "ครอบครัว"

    IsNotBuyHome = "X" if home_request.IsNotBuyHome else "  "
    IsNotOwnHome = "X" if home_request.IsNotOwnHome else "  "
    IsNeverRTAFHome = "X" if home_request.IsNeverRTAFHome else "  "
    ImportanceDuty = "X" if home_request.ImportanceDuty else "  "
    IsMoveFromOtherUnit = "X" if home_request.IsMoveFromOtherUnit else "  "
    IsHomelessEvict = "X" if home_request.IsHomelessEvict else "  "
    ContinueHouse = "X" if home_request.ContinueHouse else "  "
    IsHomelessDisaster = "X" if home_request.IsHomelessDisaster else "  "
    OtherTrouble = "X" if home_request.OtherTrouble else "  "
    IsHomeNeed = "X" if home_request.IsHomeNeed else "  "
    IsFlatNeed = "X" if home_request.IsFlatNeed else "  "
    IsShopHouseNeed = "X" if home_request.IsShopHouseNeed else "  "
    type_need = ""
    type_need = "บ้าน น." if home_request.IsHomeNeed else ""
    type_need += " แฟลต" if home_request.IsFlatNeed else ""
    type_need += " เรือนแถว" if home_request.IsShopHouseNeed else ""
    NumResidence = home_request.CoResident.all().count()
    NumResidence = "-" if NumResidence == 0 else NumResidence

    HouseRegistration = "X" if home_request.HouseRegistration else "  "
    MarriageRegistration = "X" if home_request.MarriageRegistration else "  "
    SpouseApproved = "X" if home_request.SpouseApproved else "  "
    DivorceRegistration = "X" if home_request.DivorceRegistration else "  "
    SpouseDeathRegistration = "X" if home_request.SpouseDeathRegistration else "  "

    the_day = home_request.RequesterDateSend if home_request.RequesterDateSend != None else home_request.modified
    Day = the_day.day if the_day != None else the_day.day
    Month = month_text[the_day.month] if the_day != None else month_text[the_day.day]
    Year =  str((the_day.year + 543) % 100) if the_day != None else str((the_day.day + 543) % 100) 

    PR1 = home_request.get_ZoneRequestPriority1_display() if home_request.ZoneRequestPriority1 else ""
    PR2 = home_request.get_ZoneRequestPriority2_display() if home_request.ZoneRequestPriority2 else ""
    PR3 = home_request.get_ZoneRequestPriority3_display() if home_request.ZoneRequestPriority3 else ""
    PR4 = home_request.get_ZoneRequestPriority4_display() if home_request.ZoneRequestPriority4 else ""
    PR5 = home_request.get_ZoneRequestPriority5_display() if home_request.ZoneRequestPriority5 else ""
    PR6 = home_request.get_ZoneRequestPriority6_display() if home_request.ZoneRequestPriority6 else ""

    zone_order_list = [PR1,PR2,PR3,PR4,PR5,PR6]
    zone_order_list = [zone for zone in zone_order_list if zone != ""]
    zone_order = ", ".join(zone_order_list)

    FullName = home_request.FullName
    if "ว่าที่" in FullName:
        FullName = FullName[7:]
    
    co_residence = CoResident.objects.filter(home_request = home_request).order_by("Relation")
    cr_text = ""
    for cr in co_residence:
        cr_text += str(cr.FullName) + " (" 
        cr_text += str(cr.get_Relation_display()) + ")\t" 
        cr_text += f"อาชีพ {cr.Occupation}\t"  if cr.Occupation != "-" else "\t"
        cr_text += "รายได้ {:,} บาท\n".format(cr.Salary) if cr.Salary  else "ไม่มีรายได้\n"
 

    dic = {
            'Sex': self_call,
            'Rank': home_request.Requester.Title,
            'FullName':FullName,
            'PersonID':home_request.Requester.PersonID,
            'Position':home_request.Position,
            'ADDR':home_request.Address,
            'GPC':home_request.GooglePlusCodes1,
            'Distance': str(home_request.distance),
            'UnitFN':home_request.Unit.FullName,
            'UnitSN':home_request.Unit.ShortName,
            'OfficePhone':home_request.Requester.OfficePhone,
            'MobilePhone':home_request.Requester.MobilePhone,
            'Income': "{:,}".format(Income),
            'Status':"{}".format(home_request.get_Status_display()),
            'status_description':status_description,
            'family_single':family_single,
            'type_need':type_need,
            'zone_order':zone_order,
            'SpouseName': SpouseName,
            'RequesterDateSend ': f"{Day} {Month}{Year}",
            'Z1':Z1,
            'Z2':Z2,
            'Z3':Z3,
            'WifeRentalCost':WifeRentalCost,
            'RentalCost' : RentalCost,
            'Z4':IsNotBuyHome,
            'Z5':IsNotOwnHome,
            'Z6':IsNeverRTAFHome,
            'Z7':IsNotRTAFHome,
            'Z8':ImportanceDuty,
            'TravelDescription': home_request.TravelDescription,
            'OtherTrouble': home_request.OtherTrouble,
            'co_residence': cr_text,
            'Z9':IsMoveFromOtherUnit,
            'ZA':IsHomelessEvict,
            'ZB':IsHomelessDisaster,
            'ZC':IsHomeNeed,
            'ZD':IsFlatNeed,
            'ZE':IsShopHouseNeed,
            'ZF':HouseRegistration,
            'ZG':MarriageRegistration,
            'ZH':SpouseApproved,
            'ZI':DivorceRegistration,
            'ZJ':SpouseDeathRegistration,
            'ZK':ZK,
            'ZX':ContinueHouse,
            'ZY':OtherTrouble,
            'PR1': PR1,
            'PR2': PR2,
            'PR3': PR3,
            'PR4': PR4,
            'PR5': PR5,
            'PR6': PR6,
            'Residence' : str(NumResidence),
            'Month': Month,
            'Year':Year
            }
    # print(dic)
   
    for para in document.paragraphs:
        # print('para = ',para)
        for key, value in dic.items():
            if key in para.text:
                inline = para.runs
                # print('inline = ',inline)
                # Loop added to work with runs (strings with same style)
                for i in range(len(inline)):
                    if key in inline[i].text:
                        if value: # ถ้ามีการกรอกข้อมูล
                            Value = ArabicToThai(value) if key != 'GPC' else value
                        else: # ถ้าไม่มีการกรอกข้อมูล
                            Value = ""
                        text = inline[i].text.replace(key, Value)
                        inline[i].text = text

    # Prepare document for download        
    # -----------------------------
    f = BytesIO()
    document.save(f)
    length = f.tell()
    f.seek(0)
    response = HttpResponse(
        f.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    # print('docx_title = ',docx_title)
    response['Content-Disposition'] = 'attachment; filename="{}"'.format(docx_title)
    response['Content-Length'] = length
    return response

def UnitReportDocument(request, Unit_id):
    # testdoc = static('documents/test_doc.docx')
    home_request = HomeRequest.objects.filter(Unit_id = Unit_id)

    testdoc =  os.path.join(settings.TEMPLATES[0]['DIRS'][0],'documents/house_request_unit.docx')
    docx_title= f"Unit-{Unit_id}.docx"

    document = Document(testdoc)

    Month = month_text[date.today().month]
    Year =  str((date.today().year + 543) % 100)

    dic = {
            'UnitSN': request.user.CurrentUnit.ShortName,
            'OfficePhone':request.user.OfficePhone,
            'Num_US' : "5",
            'Month': Month,
            'Year':Year
            }
    # print(dic)
    
    for para in document.paragraphs:
        # print('para = ',para)
        for key, value in dic.items():
            if key in para.text:
                inline = para.runs
                # print('inline = ',inline)
                # Loop added to work with runs (strings with same style)
                for i in range(len(inline)):
                    if key in inline[i].text:
                        if value: # ถ้ามีการกรอกข้อมูล
                            Value = ArabicToThai(value) if key != 'GPC' else value
                        else: # ถ้าไม่มีการกรอกข้อมูล
                            Value = ""
                        text = inline[i].text.replace(key, Value)
                        inline[i].text = text

    # Prepare document for download        
    # -----------------------------
    f = BytesIO()
    document.save(f)
    length = f.tell()
    f.seek(0)
    response = HttpResponse(
        f.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    # print('docx_title = ',docx_title)
    response['Content-Disposition'] = 'attachment; filename="{}"'.format(docx_title)
    response['Content-Length'] = length
    return response

def ConsentForm(request):
    testdoc =  os.path.join(settings.TEMPLATES[0]['DIRS'][0],'documents/letter_of_consent.docx')
    docx_title= f"consentform.docx"
    
    document = Document(testdoc)

    spname = request.POST.get("if_spouse_name","")
    address = request.POST.get("if_address","")

    spouse = "ภรรยา" if request.user.Sex == "ชาย" else "สามี"
    selfcall = "สามี" if request.user.Sex == "ชาย" else "ภรรยา"

    day = str(date.today().day)
    month = full_month_text[date.today().month]
    year =  str((date.today().year + 543))

    dic = {
            'spouse': spouse,            
            'selfcall': selfcall,            
            'spname': spname,            
            'address': address,            
            'username':request.user.FullName,
            'day' : day,
            'month' : month,
            'year' : year,
            }

    for para in document.paragraphs:
        # print('para = ',para)
        for key, value in dic.items():
            if key in para.text:
                inline = para.runs
                # print('inline = ',inline)
                # Loop added to work with runs (strings with same style)
                for i in range(len(inline)):
                    if key in inline[i].text:
                        if value: # ถ้ามีการกรอกข้อมูล
                            Value = ArabicToThai(value) if key != 'GPC' else value
                        else: # ถ้าไม่มีการกรอกข้อมูล
                            Value = ""
                        text = inline[i].text.replace(key, Value)
                        inline[i].text = text


    f = BytesIO()
    document.save(f)
    length = f.tell()
    f.seek(0)
    response = HttpResponse(
        f.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    # print('docx_title = ',docx_title)
    response['Content-Disposition'] = 'attachment; filename="{}"'.format(docx_title)
    response['Content-Length'] = length
    return response

@login_required 
def TestExcel(request,unit_id):
    xls_unit = Unit.objects.get(id = unit_id)

    allow_access = False
    if request.user.groups.filter(name='PERSON_UNIT_ADMIN').exists():
        if request.user.CurrentUnit == xls_unit:
            allow_access = True
            
    if request.user.groups.filter(name='PERSON_ADMIN').exists():
        allow_access = True

    if not allow_access: raise PermissionDenied()

    # testdoc = static('documents/test_doc.docx')
    testxls =  os.path.join(settings.TEMPLATES[0]['DIRS'][0],'documents/unit_report.xlsx')
    # Start by opening the spreadsheet and selecting the main sheet
    workbook = load_workbook(filename=testxls)
    xls_title= f"unit_report.xlsx"

    # sheet = workbook.active

    # Write what you want into a specific cell
    sheet_number = 1
    last_insert = 0


    for rank_level in [officer_rank, non_officer_rank]:
        for status in [status_single, status_family]:            
            queryset = HomeRequest.objects.filter(Unit = xls_unit
                                        #).filter(ProcessStep = 'US'
                                        ).filter(year_round__Year = get_current_year()
                                        ).filter(Requester__Rank__in = rank_level
                                        ).filter(Status__in = status 
                                        ).order_by("-UnitTroubleScore")                      
                                        
            sheets = workbook.sheetnames
            sheet = workbook[sheets[sheet_number]]
            sheet["A1"] = f'หน่วย {request.user.CurrentUnit.ShortName}'
            first_row = 6
            
            Rank = request.user.get_Rank_display()
            if "ว่าที่" in request.user.get_Rank_display():
                Rank = Rank[7:]

            # sheet["J36"] = f"{Rank}"
            # sheet["J37"] = f"({request.user.first_name}  {request.user.last_name})"

            max_data = queryset.count()
            for i, data in enumerate(queryset):
                # if data.Requester.Sex == "หญิง":
                #     continue
                last_row = 7
                
                if i > last_row:
                    # sheet.insert_rows(i+last_row + 1)
                    # sheet.delete_rows(150)
                    for ch in "ABCDEFGHIJKL":
                        sheet[f'{ch}{i+last_row + 1}'].font = copy(sheet[f'{ch}{i+last_row-2}'].font)
                        sheet[f'{ch}{i+last_row + 1}'].border = copy(sheet[f'{ch}{i+last_row-2}'].border)
                        sheet[f'{ch}{i+last_row + 1}'].alignment = copy(sheet[f'{ch}{i+last_row-2}'].alignment)

                    print('i+last_row = ',i+last_row  )
                    
                    
                Salary = data.Salary if data.Salary else 0
                AddSalary = data.AddSalary if data.AddSalary else 0
                RentalCost = data.RentalCost if data.RentalCost else "-"


                Income = "{:,}".format(Salary + AddSalary)
                RentalCost = "{:,}".format(RentalCost) if RentalCost != "-" else "-"

                sheet[f"A{first_row+i}"] = ArabicToThai(str(i+1))
                sheet[f"B{first_row+i}"] = data.FullName
                sheet[f"C{first_row+i}"] = data.get_Status_display()
                sheet[f"D{first_row+i}"] = ArabicToThai(Income)
                sheet[f"E{first_row+i}"] = ArabicToThai(RentalCost)
                sheet[f"G{first_row+i}"] = ArabicToThai(str(data.CoResident.all().filter(Relation = '2-CH').count()))
                sheet[f"F{first_row+i}"] = ArabicToThai(str(data.CoResident.all().exclude(Relation = '2-CH').count()))
                sheet[f"L{first_row+i}"] = ArabicToThai(data.Requester.MobilePhone)
                sheet[f"H{first_row+i}"] = "ü" if data.IsHomeNeed else ""
                sheet[f"I{first_row+i}"] = "ü" if data.IsFlatNeed else ""
                sheet[f"J{first_row+i}"] = "ü" if data.IsShopHouseNeed else ""
                sheet[f"k{first_row+i}"] = ArabicToThai(str(data.UnitTroubleScore))
                sheet[f"L{first_row+i}"] = ArabicToThai(data.Requester.OfficePhone)

                last_insert += 1
            # # เลื่อนไปชีทต่อไป
            # last_row_data = i+last_row

            # # print("i+last_row = ", i+last_row)
            # sheet.move_range("I151:J155", rows = -150 + last_row_data + 4, cols=0)

            # for row_merge in range(last_row_data + 7,last_row_data + 11):
            #     sheet.merge_cells(f"J{row_merge}:L{row_merge}")
       
            sheet_number += 1         



    f = BytesIO()
    
    workbook.save(f)
    length = f.tell()
    f.seek(0)
    response = HttpResponse(
        f.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    )
    # print('xls = ',xls_title)
    response['Content-Disposition'] = 'attachment; filename="{}"'.format(xls_title)
    response['Content-Length'] = length
    return response


@login_required 
def Excel4PersonAdmin(request):

    allow_access = False
          
    if request.user.groups.filter(name='PERSON_ADMIN').exists():
        allow_access = True

    if not allow_access: raise PermissionDenied()

    # testdoc = static('documents/test_doc.docx')
    testxls =  os.path.join(settings.TEMPLATES[0]['DIRS'][0],'documents/af_report.xlsx')
    # Start by opening the spreadsheet and selecting the main sheet
    workbook = load_workbook(filename=testxls)
    xls_title= f"RequestTotal.xlsx"

    # sheet = workbook.active
    # Write what you want into a specific cell

    last_insert = 0
    
    queryset = HomeRequest.objects.filter(year_round__Year = get_current_year()
                                 ).filter(ProcessStep__in = ['PA','PP','GH']
                                 ).order_by("-PersonDateRecieved")
                                
    sheets = workbook.sheetnames
    sheet = workbook[sheets[0]]
    first_row = 2
    
    Rank = request.user.get_Rank_display()
    if "ว่าที่" in request.user.get_Rank_display():
        Rank = Rank[7:]

    # sheet["J36"] = f"{Rank}"
    # sheet["J37"] = f"({request.user.first_name}  {request.user.last_name})"
    RentPermission_text = { 1 : "ไม่มี" , 2 : "มี-ไม่เบิก", 3 : "มี-เบิก"}
    max_data = queryset.count()
    for i, data in enumerate(queryset):
            
        Salary = data.Salary if data.Salary else 0
        AddSalary = data.AddSalary if data.AddSalary else 0
        RentalCost = data.RentalCost if data.RentalCost else "-"


        Income = "{:,}".format(Salary + AddSalary)
        RentalCost = "{:,}".format(RentalCost) if RentalCost != "-" else "-"

        home_type = "ไม่ระบุ"
        if data.IsHomeNeed : 
            home_type = "บ้านพัก"
        elif data.IsFlatNeed : 
            home_type = "แฟลต"
        elif data.IsShopHouseNeed :
            home_type = "ห้องแถว"

        sheet[f"A{first_row+i}"] = "ป." if data.Requester.Rank in non_officer_rank else "ส."
        sheet[f"B{first_row+i}"] = data.Requester.get_Rank_display()
        sheet[f"C{first_row+i}"] = data.Requester.first_name + " " + data.Requester.last_name
        sheet[f"D{first_row+i}"] = str(data.Requester.CurrentUnit.ShortName)
        sheet[f"E{first_row+i}"] = data.Requester.PersonID
        sheet[f"G{first_row+i}"] = data.Requester.MobilePhone
        sheet[f"F{first_row+i}"] = data.Requester.OfficePhone
        sheet[f"L{first_row+i}"] = ArabicToThai(data.Requester.MobilePhone)
        sheet[f"H{first_row+i}"] = RentPermission_text[data.RentPermission]
        sheet[f"I{first_row+i}"] = data.Address
        sheet[f"J{first_row+i}"] = data.get_Status_display()
        sheet[f"k{first_row+i}"] = "รายงานแล้ว" if data.IsHRISReport else "-"
        sheet[f"L{first_row+i}"] = "ข้าพเจ้าไม่เป็นผู้เบิกค่าเช่าซื้อ" if data.IsNotBuyHome else "-"
        sheet[f"M{first_row+i}"] = "ข้าพเจ้าและคู่สมรสไม่มีกรรมสิทธิ์ รัศมี 20 กม." if data.IsNotOwnHome else "-"
        sheet[f"N{first_row+i}"] = "-"
        sheet[f"O{first_row+i}"] = data.OtherTrouble if data.OtherTrouble else "-"        
        sheet[f"P{first_row+i}"] =  data.Comment if data.Comment else "-"
        sheet[f"Q{first_row+i}"] =  home_type
        sheet[f"R{first_row+i}"] =  str(data.get_ZoneRequestPriority1_display()) if data.ZoneRequestPriority1 else "-"
        sheet[f"S{first_row+i}"] =  "-"
        sheet[f"T{first_row+i}"] =  data.specificed_need if data.specificed_need else "-"
        sheet[f"U{first_row+i}"] =  data.foster_person if data.foster_person else "-"
        sheet[f"V{first_row+i}"] =  "-"
        sheet[f"W{first_row+i}"] =  "-"
        sheet[f"X{first_row+i}"] =  "มี" if data.have_document else "ไม่มี"
        sheet[f"Y{first_row+i}"] =  data.document_number or "-"
        sheet[f"Z{first_row+i}"] =  str(data.document_date) if data.document_date else "-"
        sheet[f"AA{first_row+i}"] = "-" 
        sheet[f"AB{first_row+i}"] =  str(data.UnitTroubleScore) if data.UnitTroubleScore else "-"
        sheet[f"AC{first_row+i}"] =  data.get_work_commute_display() if data.work_commute else "-"
        sheet[f"AD{first_row+i}"] =  str(data.PersonTroubleScore) if data.PersonTroubleScore else "-"
        sheet[f"AK{first_row+i}"] =  str(data.get_request_type_display()) if data.request_type else "-"
        sheet[f"AL{first_row+i}"] =  str(data.get_Status_display()) if data.Status else "-"
        sheet[f"AM{first_row+i}"] =  data.num_children # "มีรายงาน" if data.IsHRISReport else "ไม่มีรายงาน"
         
        sheet[f"AN{first_row+i}"] = str(data.Salary) if data.Salary else "-"
        sheet[f"AO{first_row+i}"] =  str(data.Requester.PlacementCommandDate) if data.Requester.PlacementCommandDate else "-"
        sheet[f"AP{first_row+1}"] =  str(data.Requester.retire_date) if data.Requester.retire_date else "-"
        sheet[f"AQ{first_row+i}"] =  "-"
        sheet[f"AR{first_row+i}"] =  "-"
        sheet[f"AS{first_row+i}"] =  "-"
        sheet[f"AT{first_row+i}"] =  "-"
        sheet[f"AU{first_row+i}"] =  "-"
        sheet[f"AV{first_row+i}"] =  "-"
        sheet[f"AW{first_row+i}"] =  "-"
        sheet[f"AX{first_row+i}"] =  "-"
        sheet[f"AY{first_row+i}"] =  "-"
        sheet[f"AZ{first_row+i}"] =  str(data.PersonDateRecieved)
        sheet[f"BA{first_row+i}"] =  str(data.get_ProcessStep_display())

        last_insert += 1
            # # เลื่อนไปชีทต่อไป
            # last_row_data = i+last_row

            # # print("i+last_row = ", i+last_row)
            # sheet.move_range("I151:J155", rows = -150 + last_row_data + 4, cols=0)

            # for row_merge in range(last_row_data + 7,last_row_data + 11):
            #     sheet.merge_cells(f"J{row_merge}:L{row_merge}")
       
            # sheet_number += 1         



    f = BytesIO()
    
    workbook.save(f)
    length = f.tell()
    f.seek(0)
    response = HttpResponse(
        f.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    )
    # print('xls = ',xls_title)
    response['Content-Disposition'] = 'attachment; filename="{}"'.format(xls_title)
    response['Content-Length'] = length
    return response

#127.0.0.1:8000/hr/lnfy/
@csrf_exempt
def line_notify(request):
    CurrentYearRound = YearRound.objects.filter(CurrentStep__in = ['RS','UP','PP'])
    year_round = CurrentYearRound[0] 

    Num_RP = Count('ProcessStep', filter = Q(ProcessStep = 'RP'))
    Num_RS = Count('ProcessStep', filter = Q(ProcessStep = 'RS'))
    Num_UP = Count('ProcessStep', filter = Q(ProcessStep = 'UP'))
    Num_US = Count('ProcessStep', filter = Q(ProcessStep = 'US'))
    Num_GH = Count('ProcessStep', filter = Q(ProcessStep = 'GH'))

    homerequest = HomeRequest.objects.filter(year_round = year_round
                                    ).filter(ProcessStep__in = ['RP','RS','UP','US','GH']   
                                    ).values('year_round'                             
                                    ).annotate(                                    
                                        Num_RP = Num_RP,
                                        Num_RS = Num_RS,                                   
                                        Num_UP = Num_UP,
                                        Num_US = Num_US,
                                        Num_GH = Num_GH,
                                        waited_doc = F('Num_RS') + F('Num_UP') + F('Num_US')                                        
                                    ).values('Num_RP','Num_RS', 'Num_UP', 'Num_US', 'Num_GH', YearRound = F('year_round__Year')
                                    ).exclude(Q(Num_RP = 0) & Q(Num_RS = 0) & Q(Num_UP = 0) & Q(Num_US = 0) & Q(Num_GH = 0)
                                    ).order_by("-Num_RS","-Num_UP","-Num_US","-waited_doc")

    data = list(homerequest)
    return JsonResponse(data, safe=False)

        
from django.conf import settings
from django.http import HttpResponse, Http404
from PyPDF2 import PdfFileMerger, PdfFileReader, PdfFileWriter

@login_required
def download_decryp(request, hr_id, evidence):
    try:
        home_request = HomeRequest.objects.get(id = hr_id)
    except:
        raise Http404

    def evd_enc(evd,hr_field):
        if hr_field:
            return evd, str(hr_field)
        else:
            raise Http404

    if evidence == 'HR':
        evidence, encryp_file = evd_enc('HouseRegistration', home_request.HouseRegistration)        
    elif evidence == 'MR':
        evidence, encryp_file = evd_enc('MarriageRegistration',home_request.MarriageRegistration)        
    elif evidence == 'SA':
        evidence, encryp_file = evd_enc('SpouseApproved',home_request.SpouseApproved)        
    elif evidence == 'DR':
        evidence, encryp_file = evd_enc('DivorceRegistration',home_request.DivorceRegistration)        
    elif evidence == 'SD':
        evidence, encryp_file = evd_enc('SpouseDeathRegistration',home_request.SpouseDeathRegistration)        
    else:
        raise Http404

    # ผู้เข้าดูเอกสารได้มี 
    #   1. เจ้าของเอกสาร
    #   2. Admin ของหน่วยผู้ส่ง 
    #   3. Admin กพ.ทอ.
    allow_access = False
    access_type = ""
    if home_request.Requester.id == request.user.id:
        access_type = "Self Access, "
        allow_access = True
    if request.user.groups.filter(name='PERSON_UNIT_ADMIN').exists():
        if request.user.CurrentUnit == home_request.Unit:
            access_type += "Unit ADMIN Access, "
            allow_access = True
    if request.user.groups.filter(name='PERSON_ADMIN').exists():
        access_type += "Person ADMIN Access "
        allow_access = True

    if not allow_access: raise PermissionDenied()

    file_path = os.path.join(settings.MEDIA_ROOT, encryp_file)
    watermark_file = os.path.join(settings.TEMPLATES[0]['DIRS'][0],'documents/watermark.pdf')

    if os.path.exists(file_path):
        with open(file_path, 'rb') as fh:
            if ".enc" in encryp_file:
                data_decryp_file = decryp_file(fh.read())
            else:
                data_decryp_file = fh.read()
            
            data_decryp_file_byte_io = BytesIO(data_decryp_file)

            if not "Self" in access_type:
                input_pdf = PdfFileReader(data_decryp_file_byte_io)
                watermark_pdf = PdfFileReader(watermark_file)
                watermark_page = watermark_pdf.getPage(0)

                output = PdfFileWriter()

                for i in range(input_pdf.getNumPages()):
                    pdf_page = input_pdf.getPage(i)
                    # pdf_page = watermark_page
                    # pdf_page.mergePage(input_pdf.getPage(i))
                    pdf_page.mergePage(watermark_page)
                    output.addPage(pdf_page)
                
                merged_file = BytesIO()
                output.write(merged_file)

                # length = merged_file.tell()
                merged_file.seek(0)

                data_decryp_file = merged_file.getvalue()

            evidence_logger.info(f'{home_request.Requester.username} {evidence} access by {request.user.username} ({access_type})')
            response = HttpResponse(data_decryp_file, content_type="application/pdf")
            response['Content-Disposition'] = 'inline; filename=' + os.path.basename(file_path)
            return response
    raise Http404
