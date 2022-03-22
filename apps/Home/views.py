import os
from datetime import date, timedelta
import json
from io import StringIO, BytesIO
import logging

from django.shortcuts import render
from django.views.generic import DetailView
from django.http import Http404
from django.contrib.auth.mixins import LoginRequiredMixin, UserPassesTestMixin
from django.http import JsonResponse
from django.http import HttpResponse, HttpResponseRedirect
from django.views.decorators.csrf import csrf_exempt

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Cm
from docx.text.paragraph import Paragraph
from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER


from .models import HomeData, HomeOwner
from apps.HomeRequest.models import HomeChange
from django.conf import settings
from apps.HomeRequest.views_documents import ArabicToThai, month_text, full_month_text
from apps.HomeRequest.forms_homechange import HomeChangeBlankForm
from apps.UserData.models import User
from apps.Payment.models import WaterPayment, RentPayment
from .serializers import HomeOwnerSerializer

logger = logging.getLogger('MainLog')

class HomeDetailView(DetailView):
    model = HomeData
    template_name = "Home/detail.html"

class HomeOwnerUserDetailView(LoginRequiredMixin, DetailView):
    model = HomeOwner
    template_name = "Home/detail_by_user.html"
    login_url = '/login' 

    def get(self, request, *args, **kwargs):

        if 'who' in self.kwargs:
            who = self.kwargs['who']
        try:
            if who == 'spouse':
                spouse = User.objects.get(id = request.user.id)
                user = User.objects.get(current_spouse_pid = spouse.PersonID)
            elif who == 'owner':
                user = User.objects.get(id = request.user.id)
            else:
                raise Http404()
        except User.DoesNotExist :
            raise Http404()

        home_data = HomeOwner.objects.filter(owner = user).order_by("-is_stay")

        self.object = home_data
        context = self.get_context_data(object=self.object)
        return self.render_to_response(context)

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        home_owner = context["object"][0].id
        rent_payments = RentPayment.objects.filter(home_owner = home_owner).order_by("-date")
        water_payments = WaterPayment.objects.filter(home_owner = home_owner).order_by("-date")
        context['rent_payments'] = rent_payments
        context['water_payments'] = water_payments
        context['home_change_form'] = HomeChangeBlankForm(
                                            home_owner = HomeOwner.objects.get(id = home_owner), 
                                            user = self.request.user)

        return context

class HomeOwnerDetailView(DetailView):
    model = HomeOwner
    template_name = "Home/hm_own_detail.html"


@csrf_exempt
def homeowner_api(request, username):
    message = 'only post method with secure code'

    if request.method == 'POST':
        body = json.loads(request.body)
        if body["key"] == "123":
            try:
                user = User.objects.get(username = username)             
                homeowner = HomeOwner.objects.filter(owner = user).filter(is_stay = True)
            except User.DoesNotExist:
                dump = json.dumps({'status': 'username not found'})            
                return HttpResponse(dump, content_type='application/json')

            if not homeowner.exists():
                dump = json.dumps({'status': 'home not found'})            
                return HttpResponse(dump, content_type='application/json')
            serializer = HomeOwnerSerializer(homeowner[0])
            return JsonResponse(serializer.data)
        else:
            message = f"wrong secure code {request.body}"

    dump = json.dumps({'status': message})            
    return HttpResponse(dump, content_type='application/json')

def insert_paragraph_after(paragraph, text=None, style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.add_run(text)
    if style is not None:
        new_para.style = style
    return new_para

def ContractFormDocument(request, home_data_id):
    # testdoc = static('documents/test_doc.docx')
    # home_data = HomeData.objects.filter(id = home_data_id)
    home_owner = HomeOwner.objects.filter(owner = request.user).filter(home_id = home_data_id).filter(is_stay = True)[0]

    # if home_request.ProcessStep == 'RP':
    #     testdoc =  os.path.join(settings.TEMPLATES[0]['DIRS'][0],'documents/house_request_data_1_page_draft.docx')
    #     docx_title= f"Draft-{home_request.Requester.AFID}.docx"
    # else:

    if home_owner.home.get_type_display() == 'ฟป.' and home_owner.home.zone == '3':
        if home_owner.home.building_number == '1':
            testdoc =  os.path.join(settings.TEMPLATES[0]['DIRS'][0],'documents/contract_form_b1.docx')
        elif home_owner.home.building_number == '17':
            testdoc =  os.path.join(settings.TEMPLATES[0]['DIRS'][0],'documents/contract_form_b17.docx')
    else:
        testdoc =  os.path.join(settings.TEMPLATES[0]['DIRS'][0],'documents/contract_form.docx')

    docx_title= f"contract_form_{home_owner.owner.username}.docx"
    document = Document(testdoc)

    self_call = "กระผม" if home_owner.owner.Sex == "ชาย" else "ดิฉัน"

    Rank = "{}".format(home_owner.owner.get_Rank_display())
    if "ว่าที่" in Rank:
        Rank = Rank[7:]
    FullName = home_owner.owner.FullName
    if "ว่าที่" in FullName:
        FullName = FullName[7:]

    AFID = home_owner.owner.AFID if home_owner.owner.AFID else "-"
    PersonID = home_owner.owner.PersonID if home_owner.owner.PersonID else "-"
    Position = home_owner.owner.Position if home_owner.owner.Position else "-"

    home_type = is_Home = is_ShopHouse = is_flat = ""    
    if home_owner.home.get_type_display() in ['พ.','น.']:
        home_type = "บ้านเดี่ยว"
        is_Home = "X"
    elif home_owner.home.get_type_display() in ['ร.','ป.','น.5']:
        home_type = "ตึกแถว"
        is_ShopHouse = "X" 
    elif "ฟ" in home_owner.home.get_type_display():
        home_type = "แฟลต"
        is_ShopHouse = "X" 

    Zone1 = "X" if home_owner.home.zone == '1'else " "
    Zone2 = "X" if home_owner.home.zone == '2'else " "
    Zone3 = "X" if home_owner.home.zone == '3'else " "
    Zone6T = "X" if home_owner.home.zone == '6T'else " "
    Zone6S = "X" if home_owner.home.zone == '6S' else " "
    
    the_day = date.today()
    Day = the_day.day if the_day != None else the_day.day
    Month = full_month_text[the_day.month] if the_day != None else full_month_text[the_day.day]
    Year =  str((the_day.year + 543)) if the_day != None else str((the_day.day + 543)) 

    day_sign = home_owner.enter_command.date_sign.day
    month_sign = month_text[home_owner.enter_command.date_sign.month]
    year_sign = (home_owner.enter_command.date_sign.year + 543) % 100
    date_sign = f"{day_sign} {month_sign}{year_sign}"

    birth_day = home_owner.owner.BirthDay.day
    birth_month = month_text[home_owner.owner.BirthDay.month]
    birth_year = (home_owner.owner.BirthDay.year + 543) % 100
    birth_date = f"{birth_day} {birth_month}{birth_year}"

    OfficePhone = MobilePhone = "-"

    if home_owner.owner.OfficePhone:
        OfficePhone = home_owner.owner.OfficePhone

    if home_owner.owner.MobilePhone:
        MobilePhone = home_owner.owner.MobilePhone

    if home_owner.owner.retire_date:
        retire_day = home_owner.owner.retire_date.day
        retire_month = month_text[home_owner.owner.retire_date.month]
        retire_year = (home_owner.owner.retire_date.year + 543) % 100
        retire_date = f"{retire_day} {retire_month}{retire_year}"
    else:
        retire_date = ""

    HomeZone = home_owner.home.get_zone_display()
    num_cat = home_owner.pet.filter(type = 'cat').count()
    num_dog = home_owner.pet.filter(type = 'dog').count()
    pet_data = "ไม่มี"
    i = 1
    if num_dog > 0:
        pet_data = f"{i}. เลี้ยงหมา จำนวน {num_dog} ตัว\t" 
        i = 2

    if num_cat > 0:
        pet_data += f"{i}. เลี้ยงแมว จำนวน {num_cat} ตัว"


    # SpouseName = home_owner.SpouseName if home_request.SpouseName else ""
    dic = {
            'Sex': self_call,
            'Rank': Rank,
            'FullName':FullName ,
            'AirforceID' : f"{AFID[0:10]}",
            'PersonID': f"{PersonID[0]}-{PersonID[1:5]}-{PersonID[5:10]}-{PersonID[10:12]}-{PersonID[12]}",
            'RTAFEmail' : home_owner.owner.username + "@rtaf.mi.th",
            'Position': Position,
            'birth_date' : birth_date,
            'retire_date' : retire_date,
            'current_status' : home_owner.owner.get_current_status_display(),
            'ADDR': "", 
            'Zone1' : Zone1, 
            'Zone2' : Zone2, 
            'Zone3' : Zone3, 
            'Zone6T' : Zone6T, 
            'Zone6S' : Zone6S, 
            'home_type' : home_type,
            'is_Home' : is_Home,
            'is_ShopHouse' : is_ShopHouse,
            'is_flat' : is_flat,
            'UnitFN': home_owner.owner.CurrentUnit.FullName,
            'UnitSN': home_owner.owner.CurrentUnit.ShortName,
            'OfficePhone': OfficePhone,
            'MobilePhone': MobilePhone,
            'Month': Month,
            'Year':Year,
            'HomeZone' : HomeZone,
            'HomeCall' : home_owner.home.__str__(),
            'EnterCommand' : f"{home_owner.enter_command.number}/{home_owner.enter_command.year}",
            'date_sign' : date_sign,
            'command_name' : home_owner.enter_command.name,
            'pet_data' : pet_data,
            'cores' : "",            
            'coresdata' : "",            
            }

    table = document.tables[0]
    # row0 = t.rows[0] # for example
    # row1 = t.rows[-1]
    # row0._tr.addnext(row1._tr)
    table.cell(1, 0).text = "1"
    table.cell(1, 1).text = dic["FullName"]
    table.cell(1, 2).text = dic["HomeZone"]
    table.cell(1, 3).text = home_owner.home.get_type_display() + " " + str(home_owner.home.building_number)
    table.cell(1, 4).text = str(home_owner.home.room_number)
    # table.cell(1, 5).text = str(home_owner.home.monthly_fee)
    table.rows[1].style = document.styles['NormalText']

    coresident_table = document.tables[1]
    cores = []
    for (i, cs) in enumerate(home_owner.CoResident.all()):
        try:
            cs_person_id = f"{cs.person_id[0]}-{cs.person_id[1:5]}-{cs.person_id[5:10]}-{cs.person_id[10:12]}-{cs.person_id[12]}",            
        except:
            cs_person_id = "....................."
        cores.append(ArabicToThai(f"\t{i+1}. {cs.full_name} เลขประจำตัวประชาชน {cs_person_id[0]}  อายุ {cs.age()} ปี  ความสัมพันธ์ {cs.get_relation_display()}"))
        
        coresident_table.cell(1 + i, 0).text = cs.full_name
        if cs.birth_day:
            birth_day = cs.birth_day.day
            birth_month = month_text[cs.birth_day.month]
            birth_year = (cs.birth_day.year + 543) % 100
            birth_date = f"{birth_day} {birth_month}{birth_year}"

            coresident_table.cell(1 + i, 1).text = ArabicToThai(birth_date)
            
        coresident_table.cell(1 + i, 2).text = cs.get_relation_display()
        coresident_table.cell(1 + i, 5).text = ArabicToThai(cs_person_id[0])
        coresident_table.cell(1 + i, 5).style = document.styles['NormalText']
        # coresident_table.cell(1 + i, 3).style = document.styles['NormalText']
        # coresident_table.rows[1 + i].style = document.styles['NormalText']


    
    # print(dic)
    
    vehical_table = document.tables[2]
    for (i, vh) in enumerate(home_owner.HomeParker.all()):
        vehical_table.cell(1 + i, 0).text = vh.get_type_display()
        vehical_table.cell(1 + i, 1).text = vh.brand
        vehical_table.cell(1 + i, 2).text = vh.color
        vehical_table.cell(1 + i, 3).text = vh.plate
        vehical_table.cell(1 + i, 4).text = vh.province
        # vehical_table.cell(1 + i, 3).style = document.styles['NormalText']
        vehical_table.rows[1 + i].style = document.styles['NormalText']
    

    for para in document.paragraphs:
        # print('para = ',para.text)
        for key, value in dic.items():
            if key in para.text:
                if key == "cores":
                    for cs in cores[::-1]:
                        new_para = insert_paragraph_after(para,cs)
                        tab_stops = new_para.paragraph_format.tab_stops
                        tab_stop = tab_stops.add_tab_stop(Cm(1.25), WD_TAB_ALIGNMENT.LEFT)
                    # tab_stop = tab_stops.add_tab_stop(Cm(10))
                else:
                    inline = para.runs                
                    for i in range(len(inline)):
                        # print('inline = ',inline[i].text)
                        if key in inline[i].text:
                            if value: # ถ้ามีการกรอกข้อมูล
                                Value = ArabicToThai(value) if key != 'GPC' else value
                            else: # ถ้าไม่มีการกรอกข้อมูล
                                Value = ""
                            text = inline[i].text.replace(key, Value)
                            inline[i].text = text
                        

    logger.info(f'{request.user.username} download contract form')
    # Prepare document for download        
    # -----------------------------
    f = BytesIO()
    document.save(f)
    length = f.tell()   
    f.seek(0)
    response = HttpResponse(
        f.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    # print('docx_title = ',docx_title)
    response['Content-Disposition'] = 'attachment; filename="{}"'.format(docx_title)
    response['Content-Length'] = length
    return response
